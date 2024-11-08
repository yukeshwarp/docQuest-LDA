import streamlit as st
import json
import redis
from utils.pdf_processing import process_pdf_task
from utils.llm_interaction import ask_question
from concurrent.futures import ThreadPoolExecutor, as_completed
import io
from docx import Document
import uuid
import tiktoken
from sklearn.decomposition import LatentDirichletAllocation
from sklearn.feature_extraction.text import CountVectorizer
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
import nltk

# Download NLTK data
nltk.download('stopwords')
nltk.download('punkt')

def count_tokens(text, model="gpt-4o"):
    encoding = tiktoken.encoding_for_model(model)
    tokens = encoding.encode(text)
    return len(tokens)

# Initialize Redis client without SSL
redis_client = redis.Redis(
    host="yuktestredis.redis.cache.windows.net",
    port=6379,
    password="VBhswgzkLiRpsHVUf4XEI2uGmidT94VhuAzCaB2tVjs="
)

# Initialize session state for session_id, chat_history, and doc_token
if "session_id" not in st.session_state:
    st.session_state.session_id = str(uuid.uuid4())  # Unique ID per user session
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "doc_token" not in st.session_state:
    st.session_state.doc_token = 0

def preprocess_text_for_lda(text):
    stop_words = set(stopwords.words("english"))
    words = word_tokenize(text.lower())
    return " ".join([word for word in words if word.isalpha() and word not in stop_words])

def extract_topics(text, num_topics=3, num_words=5):
    # Preprocess text for LDA
    processed_text = preprocess_text_for_lda(text)
    
    # Use CountVectorizer to transform text data into a document-term matrix
    vectorizer = CountVectorizer()
    doc_term_matrix = vectorizer.fit_transform([processed_text])
    
    # Perform LDA topic modeling
    lda_model = LatentDirichletAllocation(n_components=num_topics, random_state=0)
    lda_model.fit(doc_term_matrix)
    
    # Extract topics
    topics = []
    for idx, topic in enumerate(lda_model.components_):
        topic_terms = [vectorizer.get_feature_names_out()[i] for i in topic.argsort()[-num_words:][::-1]]
        topics.append(f"Topic {idx + 1}: " + ", ".join(topic_terms))
    
    return topics

def save_document_to_redis(session_id, file_name, document_data):
    redis_key = f"{session_id}:document_data:{file_name}"
    redis_client.set(redis_key, json.dumps(document_data))

def get_document_from_redis(session_id, file_name):
    redis_key = f"{session_id}:document_data:{file_name}"
    data = redis_client.get(redis_key)
    if data:
        return json.loads(data)
    return None

def retrieve_user_documents_from_redis(session_id):
    documents = {}
    for key in redis_client.keys(f"{session_id}:document_data:*"):
        file_name = key.decode().split(f"{session_id}:document_data:")[1]
        documents[file_name] = get_document_from_redis(session_id, file_name)
    return documents

def handle_question(prompt, spinner_placeholder):
    if prompt:
        try:
            documents_data = retrieve_user_documents_from_redis(st.session_state.session_id)

            with spinner_placeholder.container():
                st.markdown(
                    """
                    <header>
                    <div style="text-align: center;">
                        <div class="spinner" style="margin: 20px;">
                            <div class="bounce1"></div>
                            <div class="bounce2"></div>
                            <div class="bounce3"></div>
                        </div>
                    </div>
                    </header>
                    """,
                    unsafe_allow_html=True,
                )

                answer, tot_tokens = ask_question(documents_data, prompt, st.session_state.chat_history)

            st.session_state.chat_history.append(
                {
                    "question": prompt,
                    "answer": f"{answer}\nTotal tokens: {tot_tokens}",
                }
            )

        except Exception as e:
            st.error(f"Error processing question: {e}")
        finally:
            spinner_placeholder.empty()

def reset_session():
    st.session_state.chat_history = []
    st.session_state.doc_token = 0
    for key in redis_client.keys(f"{st.session_state.session_id}:document_data:*"):
        redis_client.delete(key)

def display_chat():
    if st.session_state.chat_history:
        for i, chat in enumerate(st.session_state.chat_history):
            user_message = f"<div style='padding:10px; border-radius:10px; margin:5px 0; text-align:right;'>{chat['question']}</div>"
            assistant_message = f"<div style='padding:10px; border-radius:10px; margin:5px 0; text-align:left;'>{chat['answer']}</div>"
            st.markdown(user_message, unsafe_allow_html=True)
            st.markdown(assistant_message, unsafe_allow_html=True)

def generate_word_document(content):
    doc = Document()
    doc.add_heading("Chat Response", 0)
    doc.add_paragraph(f"Question: {content['question']}")
    doc.add_paragraph(f"Answer: {content['answer']}")
    return doc

with st.sidebar:
    uploaded_files = st.file_uploader(
        "Upload your documents",
        type=["pdf", "docx", "xlsx", "pptx"],
        accept_multiple_files=True,
        help="Supports PDF, DOCX, XLSX, and PPTX formats.",
    )

    if uploaded_files:
        new_files = []
        for uploaded_file in uploaded_files:
            if not redis_client.exists(f"{st.session_state.session_id}:document_data:{uploaded_file.name}"):
                new_files.append(uploaded_file)
            else:
                st.info(f"{uploaded_file.name} is already uploaded.")

        if new_files:
            progress_text = st.empty()
            progress_bar = st.progress(0)
            total_files = len(new_files)

            with st.spinner("Learning about your document(s)..."):
                with ThreadPoolExecutor(max_workers=2) as executor:
                    future_to_file = {
                        executor.submit(
                            process_pdf_task, uploaded_file, first_file=(index == 0)
                        ): uploaded_file
                        for index, uploaded_file in enumerate(new_files)
                    }

                    for i, future in enumerate(as_completed(future_to_file)):
                        uploaded_file = future_to_file[future]
                        try:
                            document_data = future.result()
                            st.session_state.doc_token += count_tokens(str(document_data))
                            save_document_to_redis(st.session_state.session_id, uploaded_file.name, document_data)
                            
                            # Extract topics and display in chat
                            topics = extract_topics(str(document_data))
                            topic_summary = "; ".join([f"{topic}" for topic in topics])
                            st.session_state.chat_history.append(
                                {"question": f"Topics from {uploaded_file.name}", "answer": topic_summary}
                            )
                            st.success(f"{uploaded_file.name} processed and topics extracted!")
                            
                        except Exception as e:
                            st.error(f"Error processing {uploaded_file.name}: {e}")

                        progress_bar.progress((i + 1) / total_files)
            st.sidebar.write(f"Total document tokens: {st.session_state.doc_token}")
            progress_text.text("Processing complete.")
            progress_bar.empty()

    if retrieve_user_documents_from_redis(st.session_state.session_id):
        download_data = json.dumps(retrieve_user_documents_from_redis(st.session_state.session_id), indent=4)
        st.download_button(
            label="Download Document Analysis",
            data=download_data,
            file_name="document_analysis.json",
            mime="application/json",
        )

st.image("logoD.png", width=200)
st.title("docQuest")
st.subheader("Unveil the Essence, Compare Easily, Analyze Smartly", divider="orange")

if retrieve_user_documents_from_redis(st.session_state.session_id):
    prompt = st.chat_input("Ask me anything about your documents", key="chat_input")
    spinner_placeholder = st.empty()
    if prompt:
        handle_question(prompt, spinner_placeholder)

display_chat()
